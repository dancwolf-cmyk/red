import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches

# Function to draw schematic
def draw_moe(path):
    fig, ax = plt.subplots(figsize=(7,7))
    box_style = dict(boxstyle="round,pad=0.04", linewidth=1.8, edgecolor="black", facecolor="#1f77b4")
    arrow_style = dict(arrowstyle="->", linewidth=1.6, color="black")
    mx=0.221
    my=0.043
    # Input and router boxes
    ax.add_patch(patches.FancyBboxPatch((0.25+mx,0.75+my),0.05,0.002, **box_style))
    ax.text(0.5,0.7925,"Input x",ha="center",va="center", fontsize=12)
    ax.add_patch(patches.FancyBboxPatch((0.25+mx-0.27,0.59+my+0.01),0.6,0.002, **box_style))
    ax.text(0.5,0.6525,"Router g(x)",ha="center",va="center", fontsize=12)

    # Arrow from Input to Router
    ax.annotate("", xy=(0.5,0.68), xytext=(0.5,0.75), arrowprops=arrow_style)

    # # Experts and annotations
    expert_y = 0.66-0.25
    expert_width = 0.2
    expert_height = 0.01
    expert_positions = [0.1,0.4,0.7]
    expert_labels = ["Expert1","Expert2","Expert3"]
    alpha_labels = ["α1(x)","α2(x)","αk(x)"]
    for idx,(x,label) in enumerate(zip(expert_positions, expert_labels)):
        ax.add_patch(patches.FancyBboxPatch((x,expert_y),expert_width,expert_height, **box_style))
        ax.text(x + expert_width/2, expert_y + expert_height/2, label, ha="center", va="center", fontsize=11)
        center_x = x + expert_width/2
        ax.annotate("", xy=(center_x, expert_y + expert_height), xytext=(center_x, 0.59), arrowprops=arrow_style)
        ax.text(center_x, 0.51, alpha_labels[idx], ha="center", va="center", fontsize=10)
        ax.annotate("", xy=(0.5,0.26), xytext=(center_x, expert_y), arrowprops=dict(arrowstyle="->", linewidth=1.2, color="black"))

    # Weighted sum and output
    ax.text(0.5,0.24,"Weighted Sum",ha="center",va="center", fontsize=11)
    ax.add_patch(patches.FancyBboxPatch((0.3,0.05),0.4,0.09, **box_style))
    ax.text(0.5,0.095,"Output",ha="center",va="center", fontsize=12)
    ax.annotate("", xy=(0.5,0.14), xytext=(0.5,0.22), arrowprops=arrow_style)

    ax.set_xlim(0,1)
    ax.set_ylim(0,1)
    ax.axis("off")
    plt.savefig(path, dpi=300)
    plt.close()

def draw_lora(path):
    fig, ax = plt.subplots(figsize=(4,3))
    arrow_style = dict(arrowstyle="->", linewidth=1.6, color="black")
    my=0.03
    my2=0.05
    ax.add_patch(patches.FancyBboxPatch((0.35-0.09,0.6972+my),0.5,0.0278, boxstyle="round,pad=0.04"))
    ax.text(0.5,0.725,"Base Transformer Layer",ha="center")
    ax.annotate("", xy=(0.5,0.68-my2-0.01), xytext=(0.5,0.75-my2), arrowprops=arrow_style)
    ax.annotate("", xy=(0.5,0.68-my2-0.01-0.5+0.05), xytext=(0.5,0.75-my2-0.4), arrowprops=arrow_style)
    ax.add_patch(patches.FancyBboxPatch((0.35-0.09,0.3+0.04),0.5,0.25, boxstyle="round,pad=0.04"))
    for i,(y,label) in enumerate([(0.5," W + ΔW₁ (LoRA₁)  "),(0.4,"+ ΔW₂ (LoRA₂)  "),(0.3,"+ ΔW₃ (LoRA₃) ")]):
        ax.add_patch(patches.FancyBboxPatch((0.4-0.1,y+0.0311+my),0.4,0.0089, boxstyle="round,pad=0.04"))
        ax.text(0.5,y+0.04,label,ha="center")
    ax.add_patch(patches.FancyBboxPatch((0.45,0.0889+my),0.1,0.0111,boxstyle="round,pad=0.04"))
    ax.text(0.5,0.10,"Output",ha="center")
    ax.axis("off")
    plt.savefig(path, dpi=300, bbox_inches='tight')
    plt.close()

def draw_residual(path):
    """绘制残差专家 + 可选路由器结构。"""
    fig, ax = plt.subplots(figsize=(5, 5))
    my=0.05
    box_style = dict(boxstyle="round,pad=0.01", linewidth=1.6, edgecolor="black", facecolor="#1f77b4")
    arrow_style = dict(arrowstyle="->", linewidth=1.4, color="black")

    ax.add_patch(patches.FancyBboxPatch((0.33, 0.88+my-0.01), 0.34, 0.06, **box_style))
    ax.text(0.5, 0.92+my-0.02, "Input x", ha="center", va="center", fontsize=12)

    ax.add_patch(patches.FancyBboxPatch((0.2, 0.7+my), 0.6, 0.12, **box_style))
    ax.text(0.5, 0.76+my, "Shared Transformer", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.7+0.17), xytext=(0.5, 0.74+0.18), arrowprops=arrow_style)

    ax.text(0.5, 0.62+my, "Base logits ℓ_base(x)", ha="center", va="center", fontsize=11)
    ax.annotate("", xy=(0.5, 0.57+0.12), xytext=(0.5, 0.66+0.09), arrowprops=arrow_style)

    expert_positions = [0.1, 0.67]
    expert_labels = ["Residual Expert 1", "Residual Expert K"]
    for idx, x in enumerate(expert_positions):
        ax.add_patch(patches.FancyBboxPatch((x-0.05, 0.33), 0.35, 0.12, **box_style))
        ax.text(x + 0.14, 0.39, expert_labels[idx], ha="center", va="center", fontsize=11)
        ax.text(x + 0.14, 0.29, "(small scale α)", ha="center", va="center", fontsize=9)
        center_x = x + 0.14
        ax.annotate("", xy=(center_x, 0.33+0.15), xytext=(0.5, 0.55+0.1), arrowprops=arrow_style)
        ax.annotate("", xy=(0.5, 0.19+0.02), xytext=(center_x, 0.33-0.02), arrowprops=arrow_style)

    ax.add_patch(patches.FancyBboxPatch((0.33-0.03, 0.12+my-0.02), 0.4, 0.05, **box_style))
    ax.text(0.5, 0.17+my-0.05, "Optional linear router", ha="center", va="center", fontsize=11)
    # ax.annotate("", xy=(0.5, 0.12), xytext=(0.5, 0.21), arrowprops=arrow_style)

    ax.text(0.5, 0.08+my-0.05, "ℓ_router(x) = ℓ_base(x) + β Σ α_i(x) ℓ_i(x)", ha="center", va="center", fontsize=10)
    ax.annotate("", xy=(0.5, 0.10), xytext=(0.5, 0.14), arrowprops=arrow_style)

    ax.add_patch(patches.FancyBboxPatch((0.33-0.2, 0.0+my), 0.73, 0.05, **box_style))
    ax.text(0.5, 0.0, "Output", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.0), xytext=(0.5, 0.04), arrowprops=arrow_style)

    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    plt.savefig(path, dpi=300)
    plt.close()


paths = {
    "moe": "Fig_MoE.png",
    "lora": "Fig_LoRA.png",
    "res": "Fig_Residual.png"
}

draw_moe(paths["moe"])
draw_lora(paths["lora"])
draw_residual(paths["res"])

doc = Document()
doc.add_heading("Section 5.5 Schematic Comparison Figures", level=1)

doc.add_heading("Figure X.1: Classical Sparse MoE Layer", level=2)
doc.add_picture(paths["moe"], width=Inches(4))

doc.add_heading("Figure X.2: Multi-LoRA Adapter Structure", level=2)
doc.add_picture(paths["lora"], width=Inches(4))

doc.add_heading("Figure X.3: Residual Modular Experts with Router", level=2)
doc.add_picture(paths["res"], width=Inches(4))

output_doc = "Section_5_5_Professional_PNG_EPS.docx"
doc.save(output_doc)

output_doc
