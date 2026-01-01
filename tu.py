import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches


def draw_moe(path):
    """绘制带箭头的 MoE 结构图。"""
    fig, ax = plt.subplots(figsize=(5, 5))
    box_style = dict(boxstyle="round,pad=0.01", linewidth=1.8, edgecolor="black", facecolor="#1f77b4")
    arrow_style = dict(arrowstyle="->", linewidth=1.5, color="black")

    # 输入与路由器
    ax.add_patch(patches.FancyBboxPatch((0.3, 0.78), 0.4, 0.1, **box_style))
    ax.text(0.5, 0.82, "Input x", ha="center", va="center", fontsize=12)
    ax.add_patch(patches.FancyBboxPatch((0.25, 0.6), 0.5, 0.12, **box_style))
    ax.text(0.5, 0.66, "Router g(x)", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.6), xytext=(0.5, 0.72), arrowprops=arrow_style)

    # 3 个专家和 α 线路
    expert_positions = [0.1, 0.4, 0.7]
    expert_labels = ["Expert1", "Expert2", "Expert3"]
    alpha_labels = ["α1(x)", "α2(x)", "αk(x)"]
    for idx, x in enumerate(expert_positions):
        ax.add_patch(patches.FancyBboxPatch((x, 0.35), 0.2, 0.1, **box_style))
        ax.text(x + 0.1, 0.4, expert_labels[idx], ha="center", va="center", fontsize=11)
        center_x = x + 0.1
        ax.annotate("", xy=(center_x, 0.35), xytext=(center_x, 0.55), arrowprops=arrow_style)
        ax.text(center_x, 0.52, alpha_labels[idx], ha="center", va="center", fontsize=10)
        ax.annotate("", xy=(0.5, 0.26), xytext=(center_x, 0.35), arrowprops=arrow_style)

    # Weighted sum & 输出
    ax.text(0.5, 0.26, "Weighted Sum", ha="center", va="center", fontsize=11)
    ax.add_patch(patches.FancyBboxPatch((0.3, 0.05), 0.4, 0.1, **box_style))
    ax.text(0.5, 0.095, "Output", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.05), xytext=(0.5, 0.2), arrowprops=arrow_style)

    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    plt.savefig(path, dpi=300, bbox_inches="tight")
    plt.close()


def draw_lora(path):
    """绘制 LoRA 结构示意图。"""
    fig, ax = plt.subplots(figsize=(4, 3))
    box_style = dict(boxstyle="round,pad=0.01", linewidth=1.4, edgecolor="black", facecolor="#1f77b4")
    arrow_style = dict(arrowstyle="->", linewidth=1.4, color="black")

    ax.add_patch(patches.FancyBboxPatch((0.1, 0.6), 0.8, 0.2, **box_style))
    ax.text(0.5, 0.7, "Base Transformer Layer", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.6), xytext=(0.5, 0.75), arrowprops=arrow_style)

    adapter_positions = [0.1, 0.4, 0.7]
    adapter_labels = ["LoRA1", "LoRA2", "LoRA3"]
    for idx, x in enumerate(adapter_positions):
        ax.add_patch(patches.FancyBboxPatch((x, 0.35), 0.18, 0.08, **box_style))
        ax.text(x + 0.09, 0.39, adapter_labels[idx], ha="center", va="center", fontsize=10)
        ax.annotate("", xy=(x + 0.09, 0.35), xytext=(0.5, 0.55), arrowprops=arrow_style)

    ax.add_patch(patches.FancyBboxPatch((0.4, 0.08), 0.2, 0.1, **box_style))
    ax.text(0.5, 0.13, "Output", ha="center", va="center", fontsize=11)
    ax.annotate("", xy=(0.5, 0.08), xytext=(0.5, 0.29), arrowprops=arrow_style)

    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    plt.savefig(path, dpi=300, bbox_inches="tight")
    plt.close()


def draw_residual(path):
    """绘制残差专家 + 可选路由器结构。"""
    fig, ax = plt.subplots(figsize=(5, 5))
    box_style = dict(boxstyle="round,pad=0.01", linewidth=1.6, edgecolor="black", facecolor="#1f77b4")
    arrow_style = dict(arrowstyle="->", linewidth=1.4, color="black")

    ax.add_patch(patches.FancyBboxPatch((0.33, 0.88), 0.34, 0.08, **box_style))
    ax.text(0.5, 0.92, "Input x", ha="center", va="center", fontsize=12)

    ax.add_patch(patches.FancyBboxPatch((0.2, 0.7), 0.6, 0.12, **box_style))
    ax.text(0.5, 0.76, "Shared Transformer", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.7), xytext=(0.5, 0.74), arrowprops=arrow_style)

    ax.text(0.5, 0.62, "Base logits ℓ_base(x)", ha="center", va="center", fontsize=11)
    ax.annotate("", xy=(0.5, 0.57), xytext=(0.5, 0.66), arrowprops=arrow_style)

    expert_positions = [0.05, 0.67]
    expert_labels = ["Residual Expert 1", "Residual Expert K"]
    for idx, x in enumerate(expert_positions):
        ax.add_patch(patches.FancyBboxPatch((x, 0.33), 0.28, 0.12, **box_style))
        ax.text(x + 0.14, 0.39, expert_labels[idx], ha="center", va="center", fontsize=11)
        ax.text(x + 0.14, 0.29, "(small scale α)", ha="center", va="center", fontsize=9)
        center_x = x + 0.14
        ax.annotate("", xy=(center_x, 0.33), xytext=(0.5, 0.55), arrowprops=arrow_style)
        ax.annotate("", xy=(0.5, 0.19), xytext=(center_x, 0.33), arrowprops=arrow_style)

    ax.add_patch(patches.FancyBboxPatch((0.33, 0.12), 0.34, 0.1, **box_style))
    ax.text(0.5, 0.17, "Optional linear router", ha="center", va="center", fontsize=11)
    ax.annotate("", xy=(0.5, 0.12), xytext=(0.5, 0.21), arrowprops=arrow_style)

    ax.text(0.5, 0.08, "ℓ_router(x) = ℓ_base(x) + β Σ α_i(x) ℓ_i(x)", ha="center", va="center", fontsize=10)
    ax.annotate("", xy=(0.5, 0.06), xytext=(0.5, 0.12), arrowprops=arrow_style)

    ax.add_patch(patches.FancyBboxPatch((0.33, 0.0), 0.34, 0.1, **box_style))
    ax.text(0.5, 0.05, "Output", ha="center", va="center", fontsize=12)
    ax.annotate("", xy=(0.5, 0.0), xytext=(0.5, 0.06), arrowprops=arrow_style)

    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    plt.savefig(path, dpi=300, bbox_inches="tight")
    plt.close()


paths = {
    "moe": "Fig_MoE.png",
    "lora": "Fig_LoRA.png",
    "res": "Fig_Residual.png",
}

draw_moe(paths["moe"])
draw_lora(paths["lora"])
draw_residual(paths["res"])

doc = Document()
doc.add_heading("Section 5.5 Schematic Comparison Figures", level=1)

doc.add_heading("Figure X.1: Classical Sparse MoE Layer", level=2)
doc.add_picture(paths["moe"], width=Inches(4))

doc.add_heading("Figure X.2: Multi-LoRA Adapter Structure", level=2)
doc.add_picture(paths["lora"], width=Inches(4))

doc.add_heading("Figure X.3: Residual Modular Experts with Router", level=2)
doc.add_picture(paths["res"], width=Inches(4))

output_doc = "Section_5_5_Professional_PNG_EPS.docx"
doc.save(output_doc)

output_doc
